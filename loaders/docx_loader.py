import os
import docx
from typing import List
from llama_index.core import Document
from app.loaders.base_loader import BaseLoader
from app.loaders.text_cleaner import full_clean


class DocxLoader(BaseLoader):
    """Word (.docx) dosyalarını okuyup temizleyen yükleyici."""

    def load(self, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Dosya bulunamadı: {file_path}")

        file_name = os.path.basename(file_path)
        doc = docx.Document(file_path)

        # Boş paragrafları filtrele, ardından birleştir
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        raw_text = "\n\n".join(paragraphs)

        # Temizlik pipeline'ını uygula
        clean = full_clean(raw_text)

        if not clean:
            return []

        # Word belgelerinde sayfa kavramı sabit olmadığından sayfa no: 1 verilir
        return [
            Document(
                text=clean,
                metadata={
                    "file_name": file_name,
                    "page_number": 1,
                },
            )
        ]
