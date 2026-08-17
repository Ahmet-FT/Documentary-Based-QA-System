"""
document_loader.py
==================
Tüm dosya formatları için tek bir giriş noktası.

Desteklenen formatlar: .pdf · .docx · .txt · .md

Her Document nesnesine aşağıdaki metadata alanları eklenir:

  Kaynak bilgisi:
    source          : Dosyanın mutlak yolu (alıntı/kaynak gösterme için)
    file_name       : Sadece dosya adı (örn: "rapor.pdf")
    file_extension  : Uzantı (örn: ".pdf")
    file_size_bytes : Dosya boyutu (byte)
    modified_at     : Son değiştirilme tarihi (ISO 8601)

  İçerik bilgisi:
    loader          : Kullanılan yükleyici ("pdf" / "docx" / "txt")
    page_number     : Sayfa numarası (PDF'de gerçek sayfa; diğerlerinde 1)
    total_pages     : Toplam sayfa sayısı (PDF'de gerçek; diğerlerinde 1)
    word_count      : Bu sayfadaki/belgdeki kelime sayısı
    char_count      : Bu sayfadaki/belgdeki karakter sayısı

Kullanım:
    from app.loaders.document_loader import DocumentLoader

    loader = DocumentLoader()
    docs = loader.load("rapor.pdf")
    # veya birden fazla dosya:
    docs = loader.load_many(["a.pdf", "b.docx"])
"""

from __future__ import annotations

import os
from datetime import datetime
from typing import List

import fitz          # PyMuPDF — PDF
import docx as _docx # python-docx — DOCX
from llama_index.core import Document

from app.loaders.text_cleaner import full_clean


# ---------------------------------------------------------------------------
# Desteklenen format sabitleri
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


# ---------------------------------------------------------------------------
# Ana sınıf
# ---------------------------------------------------------------------------

class DocumentLoader:
    """
    Dosya uzantısını otomatik algılayarak ilgili ayrıştırıcıyı çalıştırır,
    temizlik uygular ve zengin metadata ekler.

    Args:
        skip_empty_pages (bool):
            Temizleme sonrası boş kalan sayfaları atla. Varsayılan: True.
    """

    def __init__(self, skip_empty_pages: bool = True):
        self.skip_empty_pages = skip_empty_pages

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def load(self, file_path: str) -> List[Document]:
        """
        Tek bir dosyayı okuyup LlamaIndex Document listesi olarak döner.

        Args:
            file_path: Okunacak dosyanın yolu.

        Returns:
            Temizlenmiş ve metadata eklenmiş Document nesneleri listesi.

        Raises:
            FileNotFoundError: Dosya bulunamazsa.
            ValueError: Desteklenmeyen format ise.
        """
        abs_path = os.path.abspath(file_path)

        if not os.path.exists(abs_path):
            raise FileNotFoundError(f"Dosya bulunamadi: {abs_path}")

        ext = os.path.splitext(abs_path)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Desteklenmeyen format: '{ext}'. "
                f"Desteklenenler: {sorted(SUPPORTED_EXTENSIONS)}"
            )

        # Ortak dosya metadata
        file_meta = self._file_metadata(abs_path)

        if ext == ".pdf":
            return self._load_pdf(abs_path, file_meta)
        elif ext == ".docx":
            return self._load_docx(abs_path, file_meta)
        else:  # .txt / .md
            return self._load_txt(abs_path, file_meta)

    def load_many(self, file_paths: List[str]) -> List[Document]:
        """
        Birden fazla dosyayı sırayla okur ve sonuçları birleştirir.

        Args:
            file_paths: Dosya yolları listesi.

        Returns:
            Tüm dosyalardan gelen Document'ların birleşik listesi.
        """
        all_docs: List[Document] = []
        for path in file_paths:
            all_docs.extend(self.load(path))
        return all_docs

    # ------------------------------------------------------------------
    # Format-bazlı yükleyiciler
    # ------------------------------------------------------------------

    def _load_pdf(self, abs_path: str, file_meta: dict) -> List[Document]:
        """
        PDF'i PyMuPDF ile sayfa sayfa okur.
        Her sayfaya tam metadata seti eklenir.
        """
        documents: List[Document] = []
        fitz_doc = fitz.open(abs_path)
        total_pages = len(fitz_doc)

        for page_idx, page in enumerate(fitz_doc):
            raw_text = page.get_text()
            clean    = full_clean(raw_text)

            if self.skip_empty_pages and not clean:
                continue

            metadata = {
                **file_meta,
                "loader":      "pdf",
                "page_number": page_idx + 1,
                "total_pages": total_pages,
                "word_count":  len(clean.split()),
                "char_count":  len(clean),
            }
            documents.append(Document(text=clean, metadata=metadata))

        fitz_doc.close()
        return documents

    def _load_docx(self, abs_path: str, file_meta: dict) -> List[Document]:
        """
        DOCX dosyasını python-docx ile okur.
        Boş paragrafları filtreler, double-newline ile birleştirir.
        """
        doc       = _docx.Document(abs_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        raw_text  = "\n\n".join(paragraphs)
        clean     = full_clean(raw_text)

        if self.skip_empty_pages and not clean:
            return []

        metadata = {
            **file_meta,
            "loader":          "docx",
            "page_number":     1,
            "total_pages":     1,
            "paragraph_count": len(paragraphs),
            "word_count":      len(clean.split()),
            "char_count":      len(clean),
        }
        return [Document(text=clean, metadata=metadata)]

    def _load_txt(self, abs_path: str, file_meta: dict) -> List[Document]:
        """
        TXT / Markdown dosyasını UTF-8 olarak okur.
        """
        with open(abs_path, "r", encoding="utf-8", errors="replace") as f:
            raw_text = f.read()

        clean = full_clean(raw_text)

        if self.skip_empty_pages and not clean:
            return []

        metadata = {
            **file_meta,
            "loader":      "txt",
            "page_number": 1,
            "total_pages": 1,
            "word_count":  len(clean.split()),
            "char_count":  len(clean),
        }
        return [Document(text=clean, metadata=metadata)]

    # ------------------------------------------------------------------
    # Yardimci
    # ------------------------------------------------------------------

    @staticmethod
    def _file_metadata(abs_path: str) -> dict:
        """Dosya sistemi bilgilerinden temel metadata sözlüğü üretir."""
        stat = os.stat(abs_path)
        return {
            "source":          abs_path,
            "file_name":       os.path.basename(abs_path),
            "file_extension":  os.path.splitext(abs_path)[1].lower(),
            "file_size_bytes": stat.st_size,
            "modified_at":     datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
        }

    # ------------------------------------------------------------------
    # Dunder
    # ------------------------------------------------------------------

    def __repr__(self) -> str:
        return f"DocumentLoader(skip_empty_pages={self.skip_empty_pages})"
